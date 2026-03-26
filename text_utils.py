"""
utils/text_utils.py  —  DeDupe ML
Combines original text similarity functions with Task 3 cross-format
PDF ↔ DOCX deep-normalisation and MinHash pipeline.
"""

import os
import re
import unicodedata
from difflib import SequenceMatcher

import fitz                              # PyMuPDF
from docx import Document as DocxDocument
from datasketch import MinHash

# ══════════════════════════════════════════════════════════════
#  ORIGINAL FUNCTIONS  (required by app1.py imports)
# ══════════════════════════════════════════════════════════════

def read_file(path: str) -> str:
    """Safely read a plain-text file. Returns '' on any error."""
    if not os.path.exists(path):
        return ""
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception:
        return ""


def normalize(text: str) -> list:
    """
    NLP ENHANCEMENT: Character Shingling (N-grams)
    Splits text into sliding windows of 3 characters.
    Example: "john" -> ['joh', 'ohn']
    """
    text = text.lower()
    text = re.sub(r"[^a-z0-9\s]", "", text)
    text = " ".join(text.split())

    if not text:
        return []
    if len(text) < 4:
        return [text]

    n = 3
    shingles = [text[i:i + n] for i in range(len(text) - n + 1)]
    return shingles


def get_minhash_obj(tokens: list, num_perm: int = 128) -> MinHash:
    """Create and return a MinHash object (used for LSH indexing)."""
    m = MinHash(num_perm=num_perm)
    for t in tokens:
        m.update(t.encode("utf-8"))
    return m


def get_text_similarity(file1_path: str, file2_path: str):
    """
    Calculates similarity using MinHash (NLP Shingles) and SequenceMatcher.
    Returns: (score: float, label: str)
    """
    text_a = read_file(file1_path)
    text_b = read_file(file2_path)

    tokens_a = normalize(text_a)
    tokens_b = normalize(text_b)

    # Edge cases: empty files
    if not tokens_a and not tokens_b:
        return 100.0, "Duplicate"
    if not tokens_a or not tokens_b:
        return 0.0, "Unique"

    # MinHash similarity (N-gram based)
    m1 = get_minhash_obj(tokens_a)
    m2 = get_minhash_obj(tokens_b)
    minhash_score = m1.jaccard(m2) * 100

    # Sequence similarity (exact character match, only for small files)
    seq_score = 0.0
    if len(text_a) < 5000 and len(text_b) < 5000:
        seq_score = SequenceMatcher(None, text_a, text_b).ratio() * 100

    # Take the higher of the two scores
    final_score = max(minhash_score, seq_score)
    label = "Duplicate" if final_score >= 70 else "Unique"

    return round(final_score, 2), label


# ══════════════════════════════════════════════════════════════
#  TASK 3 — CROSS-FORMAT DEEP NORMALISATION  (PDF ↔ DOCX)
# ══════════════════════════════════════════════════════════════

NUM_PERM     = 128
SHINGLE_SIZE = 3

# Invisible / noise Unicode characters
_NOISE_CHARS = re.compile(
    r"[\u00ad"          # soft hyphen
    r"\u200b-\u200f"    # zero-width spaces / direction marks
    r"\u2028\u2029"     # line/paragraph separators
    r"\ufeff"           # BOM
    r"\u00a0"           # non-breaking space
    r"]",
    re.UNICODE,
)

# Ligatures → plain ASCII
_LIGATURES = {
    "\ufb00": "ff", "\ufb01": "fi", "\ufb02": "fl",
    "\ufb03": "ffi", "\ufb04": "ffl",
    "\u00e6": "ae", "\u00c6": "AE",
    "\u0153": "oe", "\u0152": "OE",
}

# Smart quotes / dashes → ASCII
_TYPOGRAPHY = str.maketrans({
    "\u2018": "'", "\u2019": "'", "\u201c": '"', "\u201d": '"',
    "\u2013": "-", "\u2014": "-", "\u2026": "...",
    "\u00b7": ".",
})


def extract_text_from_pdf(path: str) -> str:
    """Extract raw text from a PDF using PyMuPDF."""
    doc   = fitz.open(path)
    pages = [page.get_text("text") for page in doc]
    doc.close()
    return "\n".join(pages)


def extract_text_from_docx(path: str) -> str:
    """Extract raw text from a DOCX file using python-docx."""
    doc   = DocxDocument(path)
    parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)

    for section in doc.sections:
        for hf in (section.header, section.footer):
            if hf is not None:
                for para in hf.paragraphs:
                    if para.text.strip():
                        parts.append(para.text)

    return "\n".join(parts)


def deep_normalize(raw_text: str) -> list:
    """
    Shared normalisation pipeline applied to text from BOTH PDF and DOCX
    before MinHash — ensures identical content produces identical tokens.

    Steps: ligature expansion → smart typography → noise removal →
           NFKC → whitespace collapse → lowercase → tokenise → strip punctuation
    """
    if not raw_text:
        return []

    text = raw_text

    for lig, rep in _LIGATURES.items():
        text = text.replace(lig, rep)

    text = text.translate(_TYPOGRAPHY)
    text = _NOISE_CHARS.sub(" ", text)
    text = unicodedata.normalize("NFKC", text)
    text = re.sub(r"[\r\n\t\f\v]+", " ", text)
    text = re.sub(r" +", " ", text).strip()
    text = text.lower()

    tokens = []
    for tok in text.split():
        tok = tok.strip(".,;:!?\"'()[]{}<>/\\|@#$%^&*_+=~`")
        if len(tok) >= 2:
            tokens.append(tok)

    return tokens


def tokens_to_minhash(tokens: list, num_perm: int = NUM_PERM) -> MinHash:
    """Build a MinHash from tokens using character n-gram shingles."""
    m = MinHash(num_perm=num_perm)
    for token in tokens:
        for i in range(max(1, len(token) - SHINGLE_SIZE + 1)):
            shingle = token[i: i + SHINGLE_SIZE]
            m.update(shingle.encode("utf-8"))
    return m


def get_minhash_for_file(path: str, num_perm: int = NUM_PERM):
    """
    Auto-detect format, extract text, deep-normalise, build MinHash.
    Returns None if the file is empty or unreadable.
    """
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".pdf":
            raw = extract_text_from_pdf(path)
        elif ext == ".docx":
            raw = extract_text_from_docx(path)
        else:
            raw = read_file(path)

        tokens = deep_normalize(raw)
        if not tokens:
            return None
        return tokens_to_minhash(tokens, num_perm)

    except Exception as e:
        print(f"[minhash] ERROR processing {path}: {e}")
        return None


def cross_format_similarity(path_a: str, path_b: str) -> float:
    """
    Jaccard similarity (0–100) between two files regardless of format.
    Uses deep_normalize so PDF ↔ DOCX comparisons are accurate.
    """
    m_a = get_minhash_for_file(path_a)
    m_b = get_minhash_for_file(path_b)
    if m_a is None or m_b is None:
        return 0.0
    return round(m_a.jaccard(m_b) * 100, 2)


# ── Quick demo ────────────────────────────────────────────────
if __name__ == "__main__":
    import sys
    if len(sys.argv) != 3:
        print("Usage: python utils/text_utils.py <file_a.pdf> <file_b.docx>")
        sys.exit(1)

    path_a, path_b = sys.argv[1], sys.argv[2]
    sim = cross_format_similarity(path_a, path_b)
    print(f"\n  Similarity: {sim:.2f}%")
    print("  ✅ DUPLICATE" if sim >= 95 else "  🟡 NEAR DUPLICATE" if sim >= 60 else "  🟢 UNIQUE")