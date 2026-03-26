"""
app1.py  —  DeDupe ML  (Flask backend)
Supports image sub-categories: mark_list · fee_receipt · general
"""
import os
import shutil
import uuid
import random
import csv
from datetime import datetime

from flask import (Flask, render_template, request, redirect,
                   url_for, session, jsonify, send_from_directory)
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
from authlib.integrations.flask_client import OAuth

# --- ML ---
from datasketch import MinHash, MinHashLSH

# --- Project ---
from config import Config
from models import db, User, UploadedFile, ResultHistory

# --- Utilities ---
from utils.text_utils  import get_text_similarity, get_minhash_obj as get_text_minhash, normalize, read_file
from utils.image_utils import (get_image_similarity,
                                get_image_similarity_by_subtype,
                                subtype_db_key,
                                is_valid_image_subtype,
                                IMAGE_SUBTYPES)
from utils.audio_utils import get_audio_minhash, get_audio_similarity
from utils.pdf_utils   import get_pdf_minhash
from utils.file_utils  import get_any_text_content
# ── TASK 2: DOCX preview support ─────────────────────────────
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

import fitz  # PyMuPDF — already a project dependency via pdf_utils
app = Flask(__name__)
app.config.from_object(Config)

# ─── Global LSH indexes (in-memory) ──────────────────────────
lsh_text  = MinHashLSH(threshold=0.4, num_perm=128)   # text/doc/other
lsh_pdf   = MinHashLSH(threshold=0.4, num_perm=128)   # pdf only
lsh_audio = MinHashLSH(threshold=0.35, num_perm=128)  # audio files
# Threshold 0.35: audio MFCC tokens are coarser than text tokens.
# A high threshold (0.6) was filtering out real duplicates before
# get_audio_similarity could even run. 0.35 lets more candidates
# through — the precise Jaccard score then decides final status.
lsh_cross = MinHashLSH(threshold=0.5, num_perm=128)   # cross-format (doc+pdf together)
# lsh_cross uses a slightly higher threshold (0.5) because textract
# extraction from doc vs pdf may produce slightly different tokens
# even for identical content (spacing, formatting artefacts).

# ─── OAuth ───────────────────────────────────────────────────
# ⚠️  SETUP REQUIRED — replace these two values with your real
#     GitHub OAuth App credentials from:
#     https://github.com/settings/developers → "OAuth Apps" → New OAuth App
#
#     Homepage URL:       http://127.0.0.1:5000
#     Authorization callback URL: http://127.0.0.1:5000/auth/github
#
# Option A (recommended): set environment variables and leave the
#   os.environ.get() calls below as-is — no hardcoded secrets.
# Option B: replace the os.environ.get() strings with your actual values.
GITHUB_CLIENT_ID     = os.environ.get("GITHUB_CLIENT_ID",     "YOUR_GITHUB_CLIENT_ID")
GITHUB_CLIENT_SECRET = os.environ.get("GITHUB_CLIENT_SECRET", "YOUR_GITHUB_CLIENT_SECRET")

oauth  = OAuth(app)
github = oauth.register(
    name='github',
    client_id=GITHUB_CLIENT_ID,
    client_secret=GITHUB_CLIENT_SECRET,
    # ── Explicit endpoints (required for authlib v1.x with GitHub) ──
    # GitHub does NOT expose an OIDC discovery URL, so we must list
    # every endpoint manually instead of using server_metadata_url.
    access_token_url='https://github.com/login/oauth/access_token',
    access_token_params=None,
    authorize_url='https://github.com/login/oauth/authorize',
    authorize_params=None,
    api_base_url='https://api.github.com/',
    # ── Tell authlib how to send credentials ──────────────────────
    # GitHub expects credentials in the POST body, not Basic-Auth header.
    client_kwargs={
        'scope': 'user:email',
        'token_endpoint_auth_method': 'client_secret_post',
    },
)

# ─── Directories ─────────────────────────────────────────────
BASE_DIR       = os.path.abspath(os.path.dirname(__file__))
UPLOAD_FOLDER  = os.path.join(BASE_DIR, 'uploads')
TICKETS_FILE   = os.path.join(BASE_DIR, 'support_tickets.csv')

app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
app.template_folder = 'templates'
app.static_folder   = 'static'

db.init_app(app)

# ─── Allowed extensions ───────────────────────────────────────
ALLOWED_EXTENSIONS = {
    "text":               {"txt", "doc", "docx", "md", "py", "html", "css", "js"},
    "image":              {"jpg", "jpeg", "png", "webp", "bmp"},
    # image subcategory aliases — same extensions, different DB bucket
    "image_mark_list":    {"jpg", "jpeg", "png", "webp", "bmp"},
    "image_fee_receipt":  {"jpg", "jpeg", "png", "webp", "bmp"},
    "image_general":      {"jpg", "jpeg", "png", "webp", "bmp"},
    "audio":              {"mp3", "wav", "ogg", "flac"},
    "pdf":                {"pdf"},
    "other":              {"csv", "json", "xml"},
}

# ─── Upload sub-folders (created on startup) ──────────────────
UPLOAD_SUBFOLDERS = [
    "text", "audio", "pdf", "other", "doc",
    "image_mark_list", "image_fee_receipt", "image_general",
    # Keep legacy "image" folder so existing files don't break
    "image",
]


def validate_file_type(filename: str, file_type: str) -> bool:
    """Ensure file extension matches the selected type/subtype."""
    if '.' not in filename:
        return False
    ext = filename.rsplit('.', 1)[-1].lower()
    # Resolve aliases
    category = "text" if file_type == "doc" else file_type
    return ext in ALLOWED_EXTENSIONS.get(category, set())


def rebuild_lsh():
    """Rebuild all LSH indexes from DB on startup."""
    print("🔄 Rebuilding LSH Index...")
    with app.app_context():
        # Text / doc / other → lsh_text
        for f in UploadedFile.query.filter(
                UploadedFile.file_type.in_(["text", "doc", "txt", "other"])).all():
            if os.path.exists(f.file_path):
                try:
                    tokens = normalize(get_any_text_content(f.file_path))
                    if tokens:
                        lsh_text.insert(f.filename, get_text_minhash(tokens))
                        lsh_cross.insert(f.filename, get_text_minhash(tokens))
                except Exception:
                    pass
        # PDF → lsh_pdf + lsh_cross
        for f in UploadedFile.query.filter_by(file_type="pdf").all():
            if os.path.exists(f.file_path):
                try:
                    m = get_pdf_minhash(f.file_path)
                    if m:
                        lsh_pdf.insert(f.filename, m)
                    # Also add PDF text to cross-format index
                    tokens = normalize(get_any_text_content(f.file_path))
                    if tokens:
                        lsh_cross.insert(f.filename + "__cross", get_text_minhash(tokens))
                except Exception:
                    pass
        # Audio → lsh_audio
        for f in UploadedFile.query.filter_by(file_type="audio").all():
            if os.path.exists(f.file_path):
                try:
                    m = get_audio_minhash(f.file_path)   # single-path → returns MinHash
                    if m:
                        lsh_audio.insert(f.filename, m)
                except Exception:
                    pass
    print("✅ LSH Index Ready.")


with app.app_context():
    db.create_all()
    for folder in UPLOAD_SUBFOLDERS:
        os.makedirs(os.path.join(UPLOAD_FOLDER, folder), exist_ok=True)
    rebuild_lsh()


# ─── Template filters ─────────────────────────────────────────
@app.template_filter('clean_filename')
def clean_filename_filter(s):
    if s and len(s) > 33 and s[32] == '_':
        return s[33:]
    return s


# ─── Static helpers ───────────────────────────────────────────
@app.route('/style.css')
def serve_css():
    return send_from_directory('static/css', 'style.css')

@app.route('/scripts.js')
def serve_js():
    return send_from_directory('static/js', 'scripts.js')


# ─── Page routes ─────────────────────────────────────────────
@app.route("/")
@app.route("/home")
def home():
    # Pass session data so logged-in users see their profile in the navbar
    return render_template(
        "home.html",
        user_name=session.get("user_name", ""),
        user_email=session.get("user_email", ""),
    )

@app.route("/signin")
@app.route("/signin.html")
def signin():
    return render_template("signin.html")

@app.route("/dashboard")
@app.route("/dashboard.html")
def dashboard():
    if "user_id" not in session:
        return redirect(url_for('signin'))
    
    user_email = session.get("user_email")
    user_name = session.get("user_name")
    
    # FIXED: Admin sees ALL files across the entire database
    if user_email == "manumaddirala@gmail.com":
        history = ResultHistory.query.order_by(ResultHistory.id.desc()).limit(10).all()
        total_files = UploadedFile.query.count()
        duplicates_found = ResultHistory.query.filter_by(status="Duplicate").count()
    # Regular user sees only their own files
    else:
        history = ResultHistory.query.filter_by(
            user_name=user_name
        ).order_by(ResultHistory.id.desc()).limit(10).all()
        total_files = UploadedFile.query.filter_by(user_id=session["user_id"]).count()
        duplicates_found = ResultHistory.query.filter_by(
            user_name=user_name, status="Duplicate"
        ).count()

    return render_template(
        "dashboard.html",
        user_name=session.get("user_name", "User"),
        user_email=user_email,
        history=history,
        total_files=total_files,
        duplicates_found=duplicates_found,
    )

@app.route("/duplicates")
@app.route("/duplicates.html")
def duplicates():
    if "user_id" not in session:
        return redirect(url_for('signin'))
    
    user_email = session.get("user_email")
    
    # FIXED: Admin sees ALL duplicate records
    if user_email == "manumaddirala@gmail.com":
        records = ResultHistory.query.order_by(ResultHistory.id.desc()).all()
    # Regular user sees only their own records
    else:
        records = ResultHistory.query.filter_by(
            user_name=session.get("user_name")
        ).order_by(ResultHistory.id.desc()).all()

    return render_template(
        "duplicates.html",
        user_name=session.get("user_name", "User"),
        user_email=user_email,
        duplicates=records,
    )

@app.route("/upload")
@app.route("/upload_files.html")
def upload():
    if "user_id" not in session:
        return redirect(url_for('signin'))
    return render_template(
        "upload_files.html",
        user_name=session.get("user_name", "User"),
        user_email=session.get("user_email", ""),
    )

@app.route("/about")
@app.route("/about.html")
def about():
    return render_template("about.html",
                           user_name=session.get("user_name", "User"),
                           user_email=session.get("user_email", ""))

@app.route("/contact")
@app.route("/contact.html")
def contact():
    return render_template("contact.html",
                           user_name=session.get("user_name", "User"),
                           user_email=session.get("user_email", ""))


# ─── OAuth ───────────────────────────────────────────────────
@app.route("/login/github")
def login_github():
    # Guard: if credentials are still placeholders, show a helpful error
    if GITHUB_CLIENT_ID == "YOUR_GITHUB_CLIENT_ID":
        return """
        <html><body style="font-family:sans-serif;padding:40px;background:#0d1220;color:#e2e8f4;">
        <h2 style="color:#f5c842;">GitHub OAuth Not Configured</h2>
        <p style="color:#94a3b8;">You need to set your GitHub OAuth credentials before this button works.</p>
        <ol style="color:#94a3b8;line-height:2;">
          <li>Go to <a href="https://github.com/settings/developers" target="_blank" style="color:#3b82f6;">github.com/settings/developers</a></li>
          <li>Click <strong style="color:#e2e8f4;">New OAuth App</strong></li>
          <li>Set Homepage URL: <code>http://127.0.0.1:5000</code></li>
          <li>Set Callback URL: <code>http://127.0.0.1:5000/auth/github</code></li>
          <li>Copy Client ID and Secret into app.py or set as environment variables:<br>
            <code style="color:#f5c842;">set GITHUB_CLIENT_ID=your_id_here</code><br>
            <code style="color:#f5c842;">set GITHUB_CLIENT_SECRET=your_secret_here</code>
          </li>
        </ol>
        <p><a href="/signin" style="color:#3b82f6;">← Back to Sign In</a></p>
        </body></html>
        """
    redirect_uri = url_for('auth_github', _external=True)
    return github.authorize_redirect(redirect_uri)

@app.route("/auth/github")
def auth_github():
    try:
        token = github.authorize_access_token()
    except Exception as e:
        # Common causes: mismatched state (CSRF), wrong credentials,
        # or the user denied access on GitHub.
        err = str(e)
        if "mismatching_state" in err or "state" in err.lower():
            return redirect(url_for('signin') + "?error=state_mismatch")
        return redirect(url_for('signin') + f"?error=oauth_failed")

    if not token:
        return redirect(url_for('signin') + "?error=no_token")

    try:
        # ── Step 1: basic profile ─────────────────────────────
        user_resp = github.get('user', token=token)
        user_info = user_resp.json()

        # ── Step 2: get verified email ────────────────────────
        # The /user endpoint only returns email if the user has set
        # it as public. We MUST call /user/emails to reliably get it.
        email = user_info.get('email')
        if not email:
            emails_resp = github.get('user/emails', token=token)
            emails = emails_resp.json()
            # Pick the primary verified email first, else first verified
            primary = next(
                (e['email'] for e in emails if e.get('primary') and e.get('verified')),
                None
            )
            verified = next(
                (e['email'] for e in emails if e.get('verified')),
                None
            )
            email = primary or verified or f"{user_info['login']}@github.user"

        # ── Step 3: find or create user ───────────────────────
        user = User.query.filter_by(email=email).first()
        if not user:
            display_name = user_info.get('name') or user_info.get('login', 'GitHub User')
            user = User(
                name=display_name,
                email=email,
                password=generate_password_hash(uuid.uuid4().hex),
            )
            db.session.add(user)
            db.session.commit()

        # ── Step 4: set session and redirect ──────────────────
        session.permanent = True
        session["user_id"]    = user.id
        session["user_name"]  = user.name
        session["user_email"] = user.email
        return redirect(url_for('dashboard'))

    except Exception as e:
        # Return a readable error page instead of a raw 400
        return f"""
        <html><body style="font-family:sans-serif;padding:40px;background:#0d1220;color:#e2e8f4;">
        <h2 style="color:#f43f5e;">GitHub OAuth Error</h2>
        <p style="color:#94a3b8;">Could not complete sign-in. Details:</p>
        <pre style="background:#141d2e;padding:16px;border-radius:8px;color:#fbbf24;">{str(e)}</pre>
        <p><a href="/signin" style="color:#3b82f6;">← Back to Sign In</a></p>
        <hr style="border-color:#1a2438;margin:24px 0;">
        <p style="color:#64748b;font-size:12px;">
        Common fixes:<br>
        1. Check GITHUB_CLIENT_ID and GITHUB_CLIENT_SECRET are correct<br>
        2. Verify the callback URL in your GitHub OAuth App is exactly: <code>http://127.0.0.1:5000/auth/github</code><br>
        3. Make sure your Flask SECRET_KEY is set and consistent
        </p>
        </body></html>
        """, 400


# ─── Auth API ────────────────────────────────────────────────
@app.route("/api/signup", methods=["POST"])
def api_signup():
    data = request.json
    name, email, password = data.get("name"), data.get("email"), data.get("password")
    if not all([name, email, password]):
        return jsonify({"success": False, "message": "All fields are required"}), 400
    if User.query.filter_by(email=email).first():
        return jsonify({"success": False, "message": "Email already registered"}), 400
    new_user = User(name=name, email=email, password=generate_password_hash(password))
    db.session.add(new_user)
    db.session.commit()
    session["user_id"]    = new_user.id
    session["user_name"]  = new_user.name
    session["user_email"] = new_user.email
    return jsonify({"success": True, "name": new_user.name})

@app.route("/api/login", methods=["POST"])
def api_login():
    data = request.json
    user = User.query.filter_by(email=data.get("email")).first()
    if user and check_password_hash(user.password, data.get("password", "")):
        session["user_id"]    = user.id
        session["user_name"]  = user.name
        session["user_email"] = user.email
        return jsonify({"success": True, "name": user.name})
    return jsonify({"success": False, "message": "Invalid credentials"}), 401


# ─── Upload API ───────────────────────────────────────────────
@app.route("/api/upload", methods=["POST"])
def api_upload():
    if "user_id" not in session:
        return jsonify({"success": False}), 401

    file        = request.files.get("file")
    file_type   = request.form.get("type", "text")          # e.g. "text", "image", "audio"
    img_subtype = request.form.get("image_subtype", "general")  # "mark_list" | "fee_receipt" | "general"

    if not file:
        return jsonify({"success": False, "message": "No file uploaded"}), 400

    # ── Resolve actual DB type key for images ──────────────────
    # When file_type == "image" we store it under its subtype bucket
    # so comparisons stay isolated per category.
    if file_type == "image":
        if not is_valid_image_subtype(img_subtype):
            img_subtype = "general"
        db_file_type = subtype_db_key(img_subtype)   # e.g. "image_mark_list"
        upload_folder_key = db_file_type             # sub-folder name matches db key
    else:
        db_file_type      = file_type
        upload_folder_key = file_type
        img_subtype       = None

    # ── Extension validation ───────────────────────────────────
    if not validate_file_type(file.filename, db_file_type):
        return jsonify({
            "success": False,
            "message": (
                f"Invalid file! You selected '{file_type}' mode"
                + (f" ({img_subtype})" if img_subtype else "")
                + " but uploaded a different file type."
            ),
        }), 400

    # ── Save file ─────────────────────────────────────────────
    unique_name = f"{uuid.uuid4().hex}_{secure_filename(file.filename)}"
    save_dir    = os.path.join(app.config["UPLOAD_FOLDER"], upload_folder_key)
    os.makedirs(save_dir, exist_ok=True)
    path = os.path.join(save_dir, unique_name)
    file.save(path)

    best_sim: float = 0.0
    matched_file: str = "None"

    # ─────────────────────────────────────────────────────────
    #  MATCHING LOGIC
    # ─────────────────────────────────────────────────────────

    # ── Text / Doc / Other ────────────────────────────────────
    if db_file_type in ["text", "doc", "txt", "other"]:
        content = get_any_text_content(path)
        tokens  = normalize(content)
        if tokens:
            m_new  = get_text_minhash(tokens)
            result = lsh_text.query(m_new)

            # Step 1: Fast LSH candidates
            for cand_key in result:
                cand = UploadedFile.query.filter_by(filename=cand_key).first()
                if cand and os.path.exists(cand.file_path):
                    try:
                        sim, _ = get_text_similarity(path, cand.file_path)
                        if sim > best_sim:
                            best_sim, matched_file = sim, cand_key
                    except Exception:
                        continue

            # Step 2: Brute-force fallback if no LSH hit
            if best_sim == 0.0:
                for f in UploadedFile.query.filter(
                        UploadedFile.file_type.in_(["text", "doc", "txt", "other"])).all():
                    if f.filename not in result and os.path.exists(f.file_path):
                        try:
                            s, _ = get_text_similarity(path, f.file_path)
                            if s > best_sim:
                                best_sim, matched_file = s, f.filename
                        except Exception:
                            continue

            if best_sim < 70:
                lsh_text.insert(unique_name, m_new)
                # Also add to cross-format index so future PDFs with same
                # content can be detected against this doc/text file
                try:
                    lsh_cross.insert(unique_name, m_new)
                except Exception:
                    pass

            # ── Cross-format: doc vs existing PDFs ────────────
            # Only for doc/docx uploads — check if same content exists as PDF
            is_cross_format = False
            cross_sim: float = 0.0
            cross_match: str = "None"
            if db_file_type in ("text", "doc") and tokens:
                try:
                    # Query lsh_cross for PDF candidates (stored with __cross suffix)
                    cross_candidates = lsh_cross.query(m_new)
                    for cand_key in cross_candidates:
                        # Strip __cross suffix to find real filename
                        real_key = cand_key.replace("__cross", "")
                        cand = UploadedFile.query.filter_by(
                            filename=real_key, file_type="pdf"
                        ).first()
                        if cand and os.path.exists(cand.file_path):
                            try:
                                cand_text   = get_any_text_content(cand.file_path)
                                cand_tokens = normalize(cand_text)
                                if cand_tokens:
                                    m_cand = get_text_minhash(cand_tokens)
                                    sim    = m_new.jaccard(m_cand) * 100
                                    if sim > cross_sim:
                                        cross_sim, cross_match = sim, real_key
                            except Exception:
                                continue
                    # Always brute-force all PDFs — reliable even when lsh_cross is empty
                    for stored in UploadedFile.query.filter_by(file_type="pdf").all():
                        if not os.path.exists(stored.file_path):
                            continue
                        try:
                            cand_text   = get_any_text_content(stored.file_path)
                            cand_tokens = normalize(cand_text)
                            if cand_tokens:
                                m_cand = get_text_minhash(cand_tokens)
                                sim    = m_new.jaccard(m_cand) * 100
                                if sim > cross_sim:
                                    cross_sim, cross_match = sim, stored.filename
                        except Exception:
                            continue
                except Exception:
                    pass

                if cross_sim >= 40 and cross_sim > best_sim:
                    best_sim        = cross_sim
                    matched_file    = cross_match
                    is_cross_format = True

    # ── PDF ───────────────────────────────────────────────────
    elif db_file_type == "pdf":
        m_new = get_pdf_minhash(path)
        if m_new:
            # Step 1: Fast LSH candidates (PDF vs PDF)
            result = lsh_pdf.query(m_new)
            for cand_key in result:
                cand = UploadedFile.query.filter_by(filename=cand_key).first()
                if cand and os.path.exists(cand.file_path):
                    try:
                        m_old = get_pdf_minhash(cand.file_path)
                        if m_old:
                            sim = m_new.jaccard(m_old) * 100
                            if sim > best_sim:
                                best_sim, matched_file = sim, cand_key
                    except Exception:
                        continue
            # Step 2: Brute-force fallback if LSH returned nothing (PDF vs PDF)
            if best_sim == 0.0:
                for stored in UploadedFile.query.filter_by(file_type="pdf").all():
                    if stored.filename in (result or []):
                        continue
                    if not os.path.exists(stored.file_path):
                        continue
                    try:
                        m_old = get_pdf_minhash(stored.file_path)
                        if m_old:
                            sim = m_new.jaccard(m_old) * 100
                            if sim > best_sim:
                                best_sim, matched_file = sim, stored.filename
                    except Exception:
                        continue

            # Step 3: Cross-format check — PDF vs DOC/DOCX/TXT
            # Uses textract to extract text from both sides and compare.
            # Threshold is 40% (not 60%) because textract produces slightly
            # different token sets from identical content in DOC vs PDF
            # (whitespace, headers, encoding differences).
            cross_sim: float  = 0.0
            cross_match: str  = "None"
            pdf_tokens        = []   # define here to avoid NameError below
            m_text            = None
            try:
                pdf_text   = get_any_text_content(path)
                pdf_tokens = normalize(pdf_text)
                if pdf_tokens:
                    m_text = get_text_minhash(pdf_tokens)

                    # Always brute-force ALL stored doc/text files
                    # (LSH may be empty on first use; brute-force is reliable)
                    all_text_files = UploadedFile.query.filter(
                        UploadedFile.file_type.in_(["text", "doc", "txt", "other"])
                    ).all()

                    for stored in all_text_files:
                        if not os.path.exists(stored.file_path):
                            continue
                        try:
                            cand_text   = get_any_text_content(stored.file_path)
                            cand_tokens = normalize(cand_text)
                            if cand_tokens:
                                m_cand = get_text_minhash(cand_tokens)
                                sim    = m_text.jaccard(m_cand) * 100
                                if sim > cross_sim:
                                    cross_sim, cross_match = sim, stored.filename
                        except Exception:
                            continue
            except Exception:
                pass

            # Cross-format threshold: 40% accounts for textract token differences
            # between same content saved as DOC vs PDF
            if cross_sim >= 40 and cross_sim > best_sim:
                best_sim        = cross_sim
                matched_file    = cross_match
                is_cross_format = True
            else:
                is_cross_format = False

            if best_sim < 95:
                try:
                    lsh_pdf.insert(unique_name, m_new)
                    # Add PDF text tokens to lsh_cross for future doc uploads
                    if pdf_tokens and m_text:
                        lsh_cross.insert(unique_name + "__cross", m_text)
                except Exception:
                    pass

    # ── Image — subcategory-isolated ──────────────────────────
    elif db_file_type.startswith("image_"):
        # Pull ALL uploaded image files; get_image_similarity_by_subtype
        # filters internally to only the matching subtype bucket.
        all_image_files = UploadedFile.query.filter(
            UploadedFile.file_type.in_([
                "image_mark_list", "image_fee_receipt", "image_general",
                "image",  # legacy bucket
            ])
        ).all()

        best_sim, matched_file = get_image_similarity_by_subtype(
            new_file_path=path,
            existing_files=all_image_files,
            image_subtype=img_subtype,
        )

    # ── Audio ─────────────────────────────────────────────────
    elif db_file_type == "audio":
        # Step 1: Get MinHash fingerprint for the new file
        m_new = get_audio_minhash(path)   # single path → returns MinHash
        if m_new:
            # Step 2: Fast LSH lookup — find candidates above threshold
            try:
                candidates = lsh_audio.query(m_new)
            except Exception:
                candidates = []

            for cand_key in candidates:
                cand = UploadedFile.query.filter_by(filename=cand_key).first()
                if cand and os.path.exists(cand.file_path):
                    try:
                        sim, _ = get_audio_similarity(path, cand.file_path)
                        if sim > best_sim:
                            best_sim, matched_file = sim, cand_key
                    except Exception:
                        continue

            # Step 3: Brute-force fallback if LSH returned nothing
            if best_sim == 0.0:
                for stored in UploadedFile.query.filter_by(file_type="audio").all():
                    if stored.filename in (candidates or []):
                        continue
                    if not os.path.exists(stored.file_path):
                        continue
                    try:
                        sim, _ = get_audio_similarity(path, stored.file_path)
                        if sim > best_sim:
                            best_sim, matched_file = sim, stored.filename
                    except Exception:
                        continue

            # Step 4: Add to LSH index if not a near-duplicate
            if best_sim < 95:
                try:
                    lsh_audio.insert(unique_name, m_new)
                except Exception:
                    pass

    # ─────────────────────────────────────────────────────────
    #  CLASSIFY & PERSIST
    # ─────────────────────────────────────────────────────────
    status = (
        "Duplicate"      if best_sim >= 70 else
        "Near Duplicate" if best_sim >= 50 else
        "Unique"
    )

    # Ensure is_cross_format is defined (image/audio paths don't set it)
    if 'is_cross_format' not in locals():
        is_cross_format = False

    # Build a human-readable display label for the subtype
    subtype_label_map = {
        "mark_list":   "Mark List",
        "fee_receipt": "Fee Receipt",
        "general":     "General Photo",
    }
    display_subtype = subtype_label_map.get(img_subtype, "") if img_subtype else ""

    # ── Cross-format match: DOC content == PDF content ────────
    # When is_cross_format is True and similarity >= 60, we do NOT
    # auto-delete. Instead we return needs_confirmation=True so the
    # UI can ask the user whether to keep both files or discard the new one.
    if is_cross_format and best_sim >= 40:
        # Save file temporarily — user will confirm/discard via /api/confirm_store
        db.session.add(UploadedFile(
            user_id   = session["user_id"],
            file_type = db_file_type,
            filename  = unique_name,
            file_path = path,
        ))
        db.session.add(ResultHistory(
            user_name     = session["user_name"],
            file_type     = db_file_type,
            uploaded_file = unique_name,
            matched_file  = matched_file,
            similarity    = round(best_sim, 2),
            status        = "Cross-Format Match",
        ))
        db.session.commit()
        return jsonify({
            "success":            True,
            "needs_confirmation": True,       # ← tells UI to show the keep/discard modal
            "cross_format":       True,
            "similarity":         round(best_sim, 2),
            "status":             "Cross-Format Match",
            "matched_file":       matched_file,
            "filename":           secure_filename(file.filename),
            "temp_filename":      unique_name,   # needed for confirm_store endpoint
            "image_subtype":      display_subtype,
        })

    db.session.add(ResultHistory(
        user_name     = session["user_name"],
        file_type     = db_file_type,
        uploaded_file = unique_name,
        matched_file  = matched_file,
        similarity    = round(best_sim, 2),
        status        = status,
    ))

    # Smart Storage: delete file if it is a pure duplicate
    if status == "Duplicate":
        try:
            os.remove(path)
            print(f"🗑️  Auto-deleted duplicate: {unique_name}")
        except Exception:
            pass
    else:
        db.session.add(UploadedFile(
            user_id   = session["user_id"],
            file_type = db_file_type,
            filename  = unique_name,
            file_path = path,
        ))

    db.session.commit()

    return jsonify({
        "success":            True,
        "needs_confirmation": False,
        "cross_format":       False,
        "similarity":         round(best_sim, 2),
        "status":             status,
        "matched_file":       matched_file,
        "filename":           secure_filename(file.filename),
        "image_subtype":      display_subtype,
    })


# ─── Delete / Reset ──────────────────────────────────────────
@app.route("/api/confirm_store", methods=["POST"])
def confirm_store():
    """
    Called by the UI after a cross-format match is detected.
    The user decides:
      action="keep"    → keep the uploaded file as-is (it stays stored)
      action="discard" → delete the uploaded file and its DB entries
    """
    if "user_id" not in session:
        return jsonify({"success": False}), 401

    data          = request.json or {}
    action        = data.get("action", "keep")       # "keep" or "discard"
    temp_filename = data.get("temp_filename", "")

    if not temp_filename:
        return jsonify({"success": False, "message": "No filename provided"}), 400

    if action == "discard":
        # Remove the file from disk and DB
        file_entry = UploadedFile.query.filter_by(filename=temp_filename).first()
        if file_entry:
            try:
                if os.path.exists(file_entry.file_path):
                    os.remove(file_entry.file_path)
            except Exception:
                pass
            db.session.delete(file_entry)
        # Update ResultHistory status to "Discarded by User"
        history = ResultHistory.query.filter_by(uploaded_file=temp_filename).first()
        if history:
            history.status = "Discarded by User"
        db.session.commit()
        return jsonify({"success": True, "action": "discarded"})

    else:  # "keep"
        # File is already saved. Just update the ResultHistory status.
        history = ResultHistory.query.filter_by(uploaded_file=temp_filename).first()
        if history:
            history.status = "Kept (Cross-Format)"
        db.session.commit()
        return jsonify({"success": True, "action": "kept"})


@app.route("/api/delete_record/<int:record_id>", methods=["DELETE"])
def delete_record(record_id):
    if "user_id" not in session:
        return jsonify({"success": False}), 401
    record = ResultHistory.query.get(record_id)
    if not record:
        return jsonify({"success": False}), 404
    file_entry = UploadedFile.query.filter_by(filename=record.uploaded_file).first()
    try:
        if file_entry and os.path.exists(file_entry.file_path):
            os.remove(file_entry.file_path)
        try:
            if record.file_type in ["text", "doc"]:
                lsh_text.remove(record.uploaded_file)
            elif record.file_type == "pdf":
                lsh_pdf.remove(record.uploaded_file)
        except Exception:
            pass
        db.session.delete(record)
        if file_entry:
            db.session.delete(file_entry)
        db.session.commit()
        return jsonify({"success": True})
    except Exception as e:
        return jsonify({"success": False, "message": str(e)}), 500


@app.route("/api/hard_reset", methods=["POST", "DELETE"])
def hard_reset():
    if "user_id" not in session:
        return jsonify({"success": False}), 401

    # ── Admin-only: only manumaddirala@gmail.com can wipe all data ──────
    if session.get("user_email") != "manumaddirala@gmail.com":
        return jsonify({"success": False, "message": "Unauthorized"}), 403

    try:
        db.session.query(ResultHistory).delete()
        db.session.query(UploadedFile).delete()
        db.session.commit()
    except Exception as e:
        return jsonify({"success": False, "message": f"DB error: {str(e)}"}), 500

    global lsh_text, lsh_pdf, lsh_audio, lsh_cross
    lsh_text  = MinHashLSH(threshold=0.4,  num_perm=128)
    lsh_pdf   = MinHashLSH(threshold=0.4,  num_perm=128)
    lsh_audio = MinHashLSH(threshold=0.35, num_perm=128)
    lsh_cross = MinHashLSH(threshold=0.5,  num_perm=128)

    for folder in UPLOAD_SUBFOLDERS:
        p = os.path.join(UPLOAD_FOLDER, folder)
        try:
            if os.path.exists(p):
                shutil.rmtree(p, ignore_errors=True)
            os.makedirs(p, exist_ok=True)
        except Exception:
            pass   # skip folders that can't be recreated (e.g. Windows file locks)

    return jsonify({"success": True})


# ─── Contact / Tickets ───────────────────────────────────────
@app.route("/api/contact", methods=["POST"])
def api_contact():
    data = request.json
    try:
        ticket_id = f"TKT-{uuid.uuid4().hex[:4].upper()}-{random.randint(100, 999)}"
        with open(TICKETS_FILE, mode='a', newline='', encoding='utf-8') as f:
            csv.writer(f).writerow([
                ticket_id,
                data.get("name"),
                data.get("email"),
                data.get("subject"),
                data.get("message"),
                datetime.now(),
            ])
        return jsonify({"success": True, "ticket_id": ticket_id, "message": "Ticket created!"})
    except Exception:
        return jsonify({"success": False}), 500


# ─── File serving ────────────────────────────────────────────
@app.route('/uploads/<file_type>/<filename>')
def serve_file(file_type, filename):
    if "user_id" not in session:
        return redirect(url_for('signin'))
    return send_from_directory(
        os.path.join(app.config["UPLOAD_FOLDER"], file_type), filename
    )
def _extract_preview_text(file_path: str) -> str:
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".docx":
        if not DOCX_AVAILABLE:
            raise ValueError("python-docx is not installed.\nRun: pip install python-docx")
        doc   = DocxDocument(file_path)
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        for section in doc.sections:
            for hf in (section.header, section.footer):
                if hf is not None:
                    for para in hf.paragraphs:
                        if para.text.strip():
                            parts.append(para.text)
        text = "\n".join(parts)
        if not text.strip():
            raise ValueError("This DOCX contains no readable text.\nIt may be image-only or protected.")
        return text

    if ext == ".pdf":
        pdf_doc = fitz.open(file_path)
        pages   = [page.get_text() for page in pdf_doc]
        pdf_doc.close()
        text = "\n".join(pages)
        if not text.strip():
            raise ValueError("This PDF is image-only — no text layer found.")
        return text

    plain_exts = {".txt", ".md", ".py", ".js", ".html", ".css", ".csv", ".json", ".xml", ".log"}
    if ext in plain_exts:
        with open(file_path, "r", encoding="utf-8", errors="replace") as fh:
            return fh.read()

    if ext == ".doc":
        if DOCX_AVAILABLE:
            try:
                doc  = DocxDocument(file_path)
                text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                if text.strip():
                    return text
            except Exception:
                pass
        raise ValueError("Legacy .doc files cannot be previewed.\nConvert to .docx or .pdf and re-upload.")

    return get_any_text_content(file_path)

@app.route('/api/preview_text/<file_type>/<filename>')
def preview_text(file_type: str, filename: str):
    if "user_id" not in session:
        return jsonify({"success": False, "error": "Unauthorized"}), 401

    # Try constructed path first
    file_path = os.path.join(app.config["UPLOAD_FOLDER"], file_type, filename)

    # DB fallback — fixes "File not found" for .docx when file_type in URL
    # doesn't match the actual subfolder the file was saved in
    if not os.path.exists(file_path):
        record = UploadedFile.query.filter_by(filename=filename).first()
        if record and os.path.exists(record.file_path):
            file_path = record.file_path
        else:
            return jsonify({
                "success": False,
                "error":   "file_not_found",
                "text":    "⚠️  File not found on disk.\n\nTry re-uploading the document.",
            }), 404

    try:
        text = _extract_preview_text(file_path)
        if not text or not text.strip():
            text = "(No readable text content found in this file.)"
        return jsonify({"success": True, "text": text})

    except ValueError as ve:
        return jsonify({
            "success": False,
            "error":   "extraction_failed",
            "text":    f"⚠️  Preview unavailable\n\n{ve}",
        }), 422

    except Exception as e:
        app.logger.exception("preview_text failed for %s / %s", file_type, filename)
        return jsonify({
            "success": False,
            "error":   "server_error",
            "text":    f"⚠️  Could not extract text.\n\nDetails: {e}",
        }), 500

# ─── Auth ────────────────────────────────────────────────────
@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for('signin'))


if __name__ == "__main__":
    app.run(debug=True, port=5001)